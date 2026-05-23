"""Generate Linq-Brand-Palette.docx — requires: pip install python-docx coloraide

Swatches prefer (1) exact hex from TS/CSS, (2) Tailwind utilities resolved from
`node_modules/tailwindcss/theme.css` OKLCH (Tailwind v4), so they match rendered UI.
Hero gradient keeps literal `#3b82f6` from `page.tsx` (vs `blue-500` token, which differs in v4).

Run from anywhere:
  pip install python-docx coloraide
  python avtive/scripts/export_brand_palette_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from coloraide import Color
except ImportError:
    print("Install coloraide: pip install coloraide", file=sys.stderr)
    raise

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROW = list[str] | tuple[list[str], list[str]]

HEX_PATTERN = re.compile(r"#([0-9A-Fa-f]{3}|[0-9A-Fa-f]{6})\b")
OKLCH_VAR = re.compile(
    r"--color-([a-z]+-\d+):\s+oklch\(([^)]+)\)\s*;",
    re.MULTILINE,
)


def expand_hex(hex_str: str) -> str:
    h = hex_str.strip().lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return h.upper()


def parse_hex_literals(*chunks: str) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for chunk in chunks:
        if not chunk:
            continue
        for m in HEX_PATTERN.findall(chunk):
            hx = expand_hex(m)
            if hx not in seen:
                seen.add(hx)
                ordered.append(hx)
    return ordered


def parse_tailwind_hex_map(theme_css_path: Path) -> dict[str, str]:
    """token e.g. 'red-500' -> RRGGBB without '#'."""
    text = theme_css_path.read_text(encoding="utf-8")
    out: dict[str, str] = {}
    for name, inner in OKLCH_VAR.findall(text):
        try:
            c = Color(f"oklch({inner.strip()})").convert("srgb")
            hx = c.to_string(hex=True).upper().lstrip("#")
            # coloraide may return 8-char with alpha sometimes; normalize to 6
            if len(hx) == 8:
                hx = hx[:6]
            out[name] = hx
        except Exception:
            continue
    return out


def tw_fill(tw: dict[str, str], *utility_keys: str) -> list[str]:
    fills: list[str] = []
    for k in utility_keys:
        if k not in tw:
            continue
        v = tw[k]
        if v not in fills:
            fills.append(v)
    return fills


def rrggbb_from_color(c: Color) -> str:
    hx = c.convert("srgb").to_string(hex=True).upper().lstrip("#")
    return hx[:6]


def css_color_mix_oklab(hex_a: str, hex_b: str, fraction_b: float) -> str:
    """Interpret as color-mix(in oklab, A (1-p)%, B p%)."""
    a = Color(hex_a if hex_a.startswith("#") else "#" + hex_a)
    b = Color(hex_b if hex_b.startswith("#") else "#" + hex_b)
    return rrggbb_from_color(a.mix(b, fraction_b, space="oklab"))


def srgb_white_over_hex(base_hex_no_hash: str, alpha_white: float) -> str:
    """Premultiplied-ish blend: foreground white with alpha atop opaque base."""
    bh = expand_hex(base_hex_no_hash)
    rb, gb_, bb = int(bh[0:2], 16), int(bh[2:4], 16), int(bh[4:6], 16)
    a = alpha_white
    r = rb * (1 - a) + 255 * a
    g = gb_ * (1 - a) + 255 * a
    b = bb * (1 - a) + 255 * a
    return f"{max(0, min(255, int(r + 0.5))):02X}{max(0, min(255, int(g + 0.5))):02X}{max(0, min(255, int(b + 0.5))):02X}"


def set_cell_shading(cell, fill_hex_6_no_hash: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill_hex_6_no_hash)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for el in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(el)
    tc_pr.append(shading)


def setup_swatch_column_outer(table, col_index: int = 0, width_inches: float = 1.05) -> None:
    for row in table.rows:
        row.cells[col_index].width = Inches(width_inches)


def fill_swatch_cell(cell, fills: list[str], max_stripes: int = 4) -> None:
    cell.text = ""
    if not fills:
        p = cell.paragraphs[0]
        p.text = "—"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(9)
        set_cell_shading(cell, "F2F2F2")
        return

    trimmed = fills[:max_stripes]
    if len(trimmed) == 1:
        set_cell_shading(cell, trimmed[0])
        cell.paragraphs[0].add_run(" ")
        return

    nested = cell.add_table(rows=1, cols=len(trimmed))
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    strip_in = max(0.09, 0.92 / len(trimmed))
    for ci, fh in enumerate(trimmed):
        c = nested.cell(0, ci)
        c.width = Inches(strip_in)
        set_cell_shading(c, fh)
        if c.paragraphs:
            c.paragraphs[0].clear()


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def row_cells(r: ROW) -> list[str]:
    return r[0] if isinstance(r, tuple) else r


def row_override(r: ROW) -> list[str] | None:
    return r[1] if isinstance(r, tuple) else None


def add_table_with_swatches(
    doc: Document,
    headers: list[str],
    rows: list[ROW],
    swatch_extractor,
) -> None:
    full_headers = ["Sample"] + headers
    tbl = doc.add_table(rows=1 + len(rows), cols=len(full_headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Sample"
    for i, h in enumerate(headers, start=1):
        hdr_cells[i].text = h
    for hp in hdr_cells[0].paragraphs + [p for hc in hdr_cells[1:] for p in hc.paragraphs]:
        for r in hp.runs:
            r.bold = True

    for ri, row in enumerate(rows):
        cells = row_cells(row)
        override = row_override(row)
        fills = override if override is not None else swatch_extractor(cells)
        dst = tbl.rows[ri + 1].cells
        fill_swatch_cell(dst[0], fills)
        for ci, val in enumerate(cells):
            dst[ci + 1].text = val

    setup_swatch_column_outer(tbl, 0)
    doc.add_paragraph()


def extractor_hex_column(hex_col: int):
    def _extract(cells: list[str]):
        if hex_col >= len(cells):
            return []
        return parse_hex_literals(cells[hex_col])

    return _extract


def extractor_all_cols(cells: list[str]):
    return parse_hex_literals(*cells)


def main():
    avtive_dir = Path(__file__).resolve().parents[1]
    repo_root = avtive_dir.parent
    out = repo_root / "Linq-Brand-Palette.docx"
    theme_css = avtive_dir / "node_modules" / "tailwindcss" / "theme.css"

    tw: dict[str, str] = {}
    if theme_css.is_file():
        tw = parse_tailwind_hex_map(theme_css)
    else:
        print(f"Warning: missing {theme_css}; Tailwind rows will use hex literals only.", file=sys.stderr)

    doc = Document()
    title = doc.add_heading("Linq brand palette (from codebase)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_para(
        doc,
        "Colors in the Sample column match the app: design tokens use hex from `globals.css`; "
        "Tailwind utilities use OKLCH from Tailwind v4 `theme.css` converted to sRGB hex; "
        "the marketing hero uses the exact arbitrary hex from `page.tsx` (including `#3b82f6`).",
    )
    doc.add_paragraph()

    # Computed matches for CSS used on the site
    body_text_hex = css_color_mix_oklab("#23468C", "#000000", 0.16)
    card_primary_end = css_color_mix_oklab("#FFFFFF", "#79D980", 0.05)
    border_overlay_10 = srgb_white_over_hex("23468C", 0.10)
    border_overlay_20 = srgb_white_over_hex("23468C", 0.20)

    slate_keys = ("slate-50", "slate-100", "slate-200", "slate-300", "slate-400", "slate-500", "slate-600")
    slate_fills = tw_fill(tw, *slate_keys)
    slate_hex_cell = (
        "Used in forms and modals: "
        + "; ".join(f"{k} #{tw[k]}" for k in slate_keys if k in tw)
    )

    red_keys = ("red-50", "red-100", "red-200", "red-500", "red-600")
    red_hex_cell = "; ".join(f"{k} #{tw[k]}" for k in red_keys if k in tw)

    amber_keys = ("amber-50", "amber-200", "amber-300", "amber-500", "amber-600", "amber-700", "amber-800")
    amber_hex_cell = "; ".join(f"{k} #{tw[k]}" for k in amber_keys if k in tw)

    green600 = tw.get("green-600", "")

    doc.add_heading("Core brand tokens (`src/app/globals.css` — `@theme`)", level=2)
    add_para(
        doc,
        "These CSS custom properties are the main design-system colors. Prefer them for new UI.",
    )
    add_table_with_swatches(
        doc,
        ["Token", "Hex", "Role"],
        [
            ["`--color-primary`", "#79D980", "Brand fill — buttons, tinted backgrounds, borders, focus rings"],
            ["`--color-primary-strong`", "#1E7A24", "Darker green for labels and strong text"],
            ["`--color-primary-foreground`", "#1C2F57", "Navy text on primary-colored buttons"],
            ["`--color-heading`", "#23468C", "Headings, hero title (`text-heading`)"],
            ["`--color-info`", "#2563EB", "Links / info (custom theme defaults in editors)"],
            ["`--color-success`", "#1E7A24", "Success semantic (aligned with primary-strong)"],
            ["`--color-warning`", "#B45309", "Defined token (amber utilities often cover UI warnings)"],
            ["`--color-danger`", "#DC2626", "Danger semantic"],
            ["`--color-surface`", "#F9FAFB", "Light shells"],
            ["`--color-border`", "#E5E7EB", "Default borders"],
            ["`--color-muted`", "#6B7280", "Secondary text"],
            ["`--color-light-1`", "#DCE7D5", "Landing radial blob (pairs with `--color-light-3`)"],
            ["`--color-light-2`", "#EDF2E9", "Landing base wash (`GradientBackground`)"],
            ["`--color-light-3`", "#DCE4F0", "Landing radial blob"],
        ],
        extractor_hex_column(1),
    )

    doc.add_heading("Tailwind mirrored colors (`tailwind.config.ts`)", level=2)
    add_para(doc, "Same hex as the `@theme` block above (`primary`, `heading`, …).")
    doc.add_paragraph()

    doc.add_heading("Layout & typography defaults", level=2)
    add_table_with_swatches(
        doc,
        ["Item", "Value", "Notes"],
        [
            (["`body` background", "solid `#FFFFFF`", "Global canvas"], parse_hex_literals("#FFFFFF")),
            (
                [
                    "`body` text color",
                    f"sRGB `#{body_text_hex}` computed",
                    "`color-mix(in oklab, heading 84%, black 16%)` (`globals.css`)",
                ],
                [body_text_hex],
            ),
            (["`h1`–`h6`", "`var(--color-heading)`", "Matches `#23468C`"], parse_hex_literals("#23468C")),
            (
                ["Muted helpers", "`--color-muted` + mixes", "`.text-muted`, `.ui-meta`, `.ui-body`"],
                [expand_hex("#6B7280")],
            ),
        ],
        extractor_hex_column(1),
    )

    doc.add_heading("Blue / green marketing accent line (`page.tsx`)", level=2)
    add_para(
        doc,
        "Hardcoded Tailwind arbitrary values — not `blue-500` token (`blue-500` in Tailwind v4 would render differently).",
    )
    add_table_with_swatches(
        doc,
        ["Gradient stops (hero subtitle)"],
        [
            (
                ["`#23468C` · `#3b82f6` · `#79D980` (as authored)"],
                ["23468C", "3B82F6", "79D980"],
            )
        ],
        extractor_hex_column(0),
    )

    doc.add_heading("Ambient backgrounds & shells", level=2)
    add_table_with_swatches(
        doc,
        ["Description", "Role"],
        [
            ([f"`light-2` `#EDF2E9`", "Full-page landing wash"], parse_hex_literals("#EDF2E9")),
            (
                ["Radial blobs `#DCE4F0` + `#DCE7D5`", "`GradientBackground.tsx` blurred glow tones"],
                ["DCE4F0", "DCE7D5"],
            ),
            (["Radial utility `#F9FAFB` → `#EBF5FF`", "`bg-radial-glow`"], parse_hex_literals("#F9FAFB", "#EBF5FF")),
            (
                ["`glass-panel` base", "~white translucent on light UI"],
                ["FFFFFF"],
            ),
            (
                [
                    "`card-primary` gradient end",
                    "`color-mix(oklab, primary 5%, white)` → shown as solid",
                ],
                ["FFFFFF", card_primary_end],
            ),
        ],
        extractor_all_cols,
    )

    doc.add_heading("Dashboard / admin chrome (`#23468C` bar + borders)", level=2)
    add_table_with_swatches(
        doc,
        ["Description", "Role"],
        [
            (["Middle stop `#2B4F95`", "`bg-linear-to-r` bars"], parse_hex_literals("#2B4F95")),
            (["Heading fill `#FFFFFF`", "Admin titles"], parse_hex_literals("#FFFFFF")),
            (["Panel end `#1e293b`", "Dark preview cards"], parse_hex_literals("#1e293b")),
            (
                [
                    "`rgba(255,255,255,0.10)` on `#23468C` → sRGB overlay",
                    f"`#{border_overlay_10}` approximate `white/10` on heading",
                ],
                [border_overlay_10],
            ),
            (
                [
                    "`rgba(255,255,255,0.20)` on `#23468C` → sRGB overlay",
                    f"`#{border_overlay_20}` approximate `white/20`",
                ],
                [border_overlay_20],
            ),
        ],
        extractor_all_cols,
    )

    doc.add_heading("Preset card themes (gradient backs)", level=2)
    add_para(doc, "`CardPreview.tsx` + editors — hex values authored in TS.")
    add_table_with_swatches(
        doc,
        ["Theme", "Gradient (start → end)", "Notes"],
        [
            (
                ["Purple", "#41295a → #2f0743", "Accent `#FFD400` + white titles (see accents column implicit)"],
                parse_hex_literals("#41295a", "#2f0743", "#FFD400"),
            ),
            (["Red", "#c94b4b → #4b134f", "White titles"], parse_hex_literals("#c94b4b", "#4b134f")),
            (["Pink", "#EE0979 → #FF6A00", "White titles"], parse_hex_literals("#EE0979", "#FF6A00")),
            (
                ["Blue", "#D3CCE3 → #E9E4F0", "Title `#5A2ED3`"],
                parse_hex_literals("#D3CCE3", "#E9E4F0", "#5A2ED3"),
            ),
        ],
        extractor_hex_column(1),
    )

    doc.add_heading("Light brand backdrops (editor previews)", level=2)
    add_table_with_swatches(
        doc,
        ["Theme", "Start → end"],
        [
            (["Purple", "#eef0ff → #f7f3ff"], parse_hex_literals("#eef0ff", "#f7f3ff")),
            (["Red", "#fff1f1 → #fdf2ff"], parse_hex_literals("#fff1f1", "#fdf2ff")),
            (["Pink", "#fff3f8 → #fff8f0"], parse_hex_literals("#fff3f8", "#fff8f0")),
            (["Blue", "#f1f5ff → #f6f8ff"], parse_hex_literals("#f1f5ff", "#f6f8ff")),
        ],
        extractor_hex_column(1),
    )

    doc.add_heading("Neutrals & Tailwind palettes (exact `theme.css` → sRGB hex)", level=2)
    add_table_with_swatches(
        doc,
        ["Use", "Hex (from Tailwind tokens)"],
        [
            ([f"White `#FFFFFF`", "`bg-white`, cards"], parse_hex_literals("#FFFFFF")),
            ([f"Scrollbar thumbs", "`#CBD5E1`, `#94A3B8` in `globals.css`"], parse_hex_literals("#CBD5E1", "#94A3B8")),
            ([slate_hex_cell, "Forms, disables, cropped modals, meta text"], slate_fills if slate_fills else []),
            (
                ["`text-red-*` / borders", red_hex_cell or "(theme.css missing)"],
                tw_fill(tw, *red_keys) if tw else [],
            ),
            (
                ["`amber-*` warnings", amber_hex_cell or "(theme.css missing)"],
                tw_fill(tw, *amber_keys) if tw else [],
            ),
            (
                ["`text-green-600` (approved)", f"`#{green600}`" if green600 else "missing theme"],
                [green600] if green600 else [],
            ),
        ],
        extractor_all_cols,
    )

    literals_block: list[ROW] = [
        (["Card frame `#141414`", "`CardPreview`"], parse_hex_literals("#141414")),
        (["Popover `#101217`", "`CustomColorPicker`"], parse_hex_literals("#101217")),
        ([f"Picker typography `#1f2937` / `#090a0c`", "`CustomColorPicker` buttons"], parse_hex_literals("#1f2937", "#090a0c")),
        ([f"Placeholder gradient `#eceff3` `#dbe3ec`", "Empty card backdrop"], parse_hex_literals("#eceff3", "#dbe3ec")),
    ]
    add_table_with_swatches(doc, ["Literal / component", "Role"], literals_block, extractor_all_cols)

    doc.add_heading("Specialty / form preview", level=2)
    add_table_with_swatches(
        doc,
        ["Hex", "Role"],
        [
            (["`#4FAE62` required chip", "`dashboard/events/[id]/page.tsx`"], parse_hex_literals("#4FAE62")),
            (["`#2F4C97` label accent", "Guest form preview"], parse_hex_literals("#2F4C97")),
        ],
        extractor_hex_column(0),
    )

    doc.add_heading("Picker default swatches", level=2)
    add_para(doc, "From `CustomColorPicker.tsx`; swatches parsed from source string.")
    add_table_with_swatches(
        doc,
        ["Description", "Colors"],
        [
            (
                ["Default greys / red presets", "`#C71B1B` + grey ramp"],
                parse_hex_literals("#C71B1B", "#F3F4F6", "#D1D5DB", "#9CA3AF", "#4B5563"),
            )
        ],
        extractor_hex_column(1),
    )

    hero_mid_hardcoded = "3B82F6"
    blue500_token = tw.get("blue-500")
    hero_note = (
        "Hero uses hardcoded `#3b82f6`. "
        + (f"Tailwind token `blue-500` resolves to `#{blue500_token}`. " if blue500_token else "")
        + "Use the Hero row for screenshots of the homepage."
    )

    doc.add_heading("Suggested export set (marketing vs tokens)", level=2)
    add_para(doc, hero_note)
    minimal_rows: list[ROW] = [
        ["Primary green", "#79D980"],
        ["Primary green dark", "#1E7A24"],
        ["Navy heading", "#23468C"],
        ["On-primary text", "#1C2F57"],
        ["Info blue (editors)", "#2563EB"],
        ["Surface / border / muted", "`#F9FAFB`, `#E5E7EB`, `#6B7280`"],
        ["Warning token / danger token", "`#B45309`, `#DC2626`"],
        (
            ["Hero mid-blue (homepage only)", "`#3b82f6`"],
            ["3B82F6"],
        ),
        ["Header chrome mid", "`#2B4F95`"],
    ]
    if blue500_token and blue500_token != hero_mid_hardcoded:
        minimal_rows.insert(
            8,
            (
                ["Tailwind `blue-500` (not hero)", f"`#{blue500_token}`"],
                [blue500_token],
            ),
        )
    add_table_with_swatches(doc, ["Role", "Hex"], minimal_rows, extractor_hex_column(1))

    doc.add_paragraph()
    add_para(
        doc,
        "Regenerate after dependency updates: Tailwind palette OKLCH can shift between minor Tailwind releases.",
        italic=True,
    )

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
