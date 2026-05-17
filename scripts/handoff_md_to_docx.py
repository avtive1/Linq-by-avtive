"""
One-off: convert docs/AUTHENTICATION-HANDOFF.md to docs/AUTHENTICATION-HANDOFF.docx
Requires: pip install python-docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


def add_hyperlink(paragraph, text: str, url: str):
    """Append a hyperlink run to paragraph (Word internal relationship)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|https?://[^\s\)\]>]+)")


def add_inline_runs(paragraph, text: str) -> None:
    """Append runs to paragraph: **bold**, `code`, URLs as hyperlinks."""
    pos = 0
    for m in _TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        chunk = m.group(1)
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif chunk.startswith("http"):
            add_hyperlink(paragraph, chunk, chunk)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def paragraph_with_inline(doc: Document, text: str) -> None:
    """Body paragraph with **bold**, `code`, and bare URLs as hyperlinks."""
    if not text.strip():
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text)


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").replace(" ", "")
    return bool(re.fullmatch(r"[\-|:]+", inner))


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            # Preserve **bold** in cells as plain for simplicity
            p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text))
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def parse_md_to_docx(md_path: Path, out_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Markdown table block
        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if is_table_sep(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                block.append(cells)
                i += 1
            add_table(doc, block)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Bullet list (skip lines that are horizontal rules made of dashes only)
        if (stripped.startswith("- ") or stripped.startswith("* ")) and not stripped.startswith(
            "---"
        ):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, text)
            i += 1
            continue

        # Numbered / checklist
        num_m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_m:
            text = num_m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, text)
            i += 1
            continue

        # Italic-only line (footer note)
        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            inner = stripped[1:-1]
            p = doc.add_paragraph()
            r = p.add_run(inner)
            r.italic = True
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        paragraph_with_inline(doc, stripped)
        i += 1

    doc.save(out_path)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "AUTHENTICATION-HANDOFF.md"
    out = root / "docs" / "AUTHENTICATION-HANDOFF.docx"
    if not md.is_file():
        print(f"Missing: {md}", file=sys.stderr)
        return 1
    parse_md_to_docx(md, out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
